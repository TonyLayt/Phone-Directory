# -*- coding: utf-8 -*-
"""
Created on Wed Dec 28 16:46:13 2022

@author: m.s.kuznietsov
"""
import sys
from PyQt5 import QtCore, QtGui, QtWidgets, QtSql
from PyQt5.QtWidgets import QMessageBox
from PyQt5.QtCore import QTimer
from uitabl import Ui_MainWindow
import json
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm
import os, time
import random
import string


class TableBook (QtWidgets.QMainWindow):
    def __init__(self):
        self.nambs = []
        QtWidgets.QMainWindow.__init__(self)
        self.table = Ui_MainWindow()
        self.table.setupUi(self)
        self.table.tableWidget.setSelectionMode(QtWidgets.QAbstractItemView.SingleSelection)
        self.get_date()
        self.update_data_list = []
        self.testlist = []
        self.dinamic_namb = []
        self.boost_row = []
        self.check = True
        self.show()
        self.time = QTimer()
        self.msg = QMessageBox()
        self.msg.setIcon(QMessageBox.Warning)        
        self.table.label_4.setText(f'<h1 style="color: rgb(73, 99, 250);">{self.date_edit}</h1>')
        self.table.label_4.setFont(QtGui.QFont("Times", 4, QtGui.QFont.Bold))
    
    def get_date (self):
        try:
            self.created = os.path.getmtime("datetable.db")         
            self.year,self.month,self.day,self.hour,self.minute,self.second = time.localtime(self.created)[:-3]
            self.date_edit = ("Дата оновлення: %02d/%02d/%d"%(self.day,self.month,self.year))
        except:
            self.date_edit = 'Записи відсутні'
    
    def Sql (self):
        self.con = QtSql.QSqlDatabase.addDatabase ('QSQLITE')
        self.con.setDatabaseName ('datetable.db')
        self.con.open()

        self.query = QtSql.QSqlQuery()
        
        if 'date' not in self.con.tables():            
            self.query.exec("""CREATE TABLE date
                        (id_1 INTEGER,
                        name TEXT,
                        position TEXT,
                        miniATC TEXT, 
                        official_namber TEXT,
                        ATC_10 TEXT,
                        office TEXT,
                        name_structure TEXT,                     
                        targetID INTEGER,
                        Booleanegt INTEGER,
                        Boolbutns INTEGER,
                        namefile)""")
        self.con.close()
        self.load_on_display()
        self.posButnonDisplay()
    
    def buttons (self):     
        self.table.lineEdit.textChanged.connect(self.find_item)
        self.table.pushButton.clicked.connect(self.slideLeftMenu)
        self.table.pushButton_2.clicked.connect(self.create_row_structure)
        self.table.pushButton_3.clicked.connect(self.create_row_persondat)
        self.table.pushButton_4.clicked.connect(self.start_server)
        self.table.pushButton_5.clicked.connect(self.edit)
        self.table.pushButton_6.clicked.connect(self.message_del)
        self.table.pushButton_7.clicked.connect(self.save_word)
        self.table.pushButton_9.clicked.connect(self.addTitleInfo)
        self.table.pushButton_10.clicked.connect(self.arrow_left)
        self.table.pushButton_11.clicked.connect(self.arrow_right)
        self.table.pushButton_12.clicked.connect(self.add_info)
        
    def create_row_structure (self):
        self.table.pushButton_6.clicked.disconnect(self.message_del)
        self.table.pushButton_6.clicked.connect(self.delete_row)
        self.table.pushButton_8.clicked.connect(self.save_structure_row)
        self.table.pushButton_12.clicked.disconnect(self.add_info)
        self.table.pushButton_12.clicked.connect(self.worningunfo)
        Id_continuous_srow = self.table.tableWidget.currentIndex().row()+1
        self.boost_row.insert(0, Id_continuous_srow)
        self.table.tableWidget.insertRow(Id_continuous_srow)
        self.table.tableWidget.setSpan (Id_continuous_srow, 0, 1, 6)
        
        styl_text_structure = QtWidgets.QTableWidgetItem()
        styl_text_structure.setTextAlignment(QtCore.Qt.AlignCenter)
        
        font_text_structure = QtGui.QFont()
        font_text_structure.setPointSize(12)
        font_text_structure.setBold(True)
        font_text_structure.setWeight(75)
        styl_text_structure.setFont(font_text_structure)
        self.table.tableWidget.setItem(Id_continuous_srow, 0, styl_text_structure)
        
        self.table.pushButton_2.blockSignals(True)
        self.table.pushButton_3.blockSignals(True)
    
    def add_row_stru_on_display (self, index, item): 
        self.table.tableWidget.insertRow(index)
        self.table.tableWidget.setSpan (index, 0, 1, 6)
              
        styl_text_load_structure = QtWidgets.QTableWidgetItem()
        styl_text_load_structure.setTextAlignment(QtCore.Qt.AlignCenter)
        
        font_text_load_structure = QtGui.QFont()
        font_text_load_structure.setPointSize(12)
        font_text_load_structure.setBold(True)
        font_text_load_structure.setWeight(75)
        styl_text_load_structure.setFont(font_text_load_structure)
        self.table.tableWidget.setItem(index, 0, styl_text_load_structure)
        item_header = self.table.tableWidget.item(index, 0)
        item_header.setText(item)
        self.table.tableWidget.resizeRowToContents(index)
        
    def create_row_persondat (self):
        self.table.pushButton_6.clicked.disconnect(self.message_del)
        self.table.pushButton_6.clicked.connect(self.delete_row)        
        self.table.pushButton_8.clicked.connect(self.save_persondat_row)
        self.table.pushButton_12.clicked.disconnect(self.add_info)
        self.table.pushButton_12.clicked.connect(self.worningunfo)
        Id_continuous_prow = self.table.tableWidget.currentIndex().row()+1
        self.boost_row.insert(0, Id_continuous_prow)
        self.table.tableWidget.insertRow(Id_continuous_prow)
       
        self.table.pushButton_2.blockSignals(True)
        self.table.pushButton_3.blockSignals(True)
    
    def add_row_persondat_on_display (self, index, item_0, item_1, item_2, item_3, item_4, item_5): 
        self.table.tableWidget.insertRow(index)    
        self.table.tableWidget.setItem(index, 0, QtWidgets.QTableWidgetItem(item_0))
        self.table.tableWidget.setItem(index, 1, QtWidgets.QTableWidgetItem(item_1))
        self.table.tableWidget.setItem(index, 2, QtWidgets.QTableWidgetItem(item_2))
        self.table.tableWidget.setItem(index, 3, QtWidgets.QTableWidgetItem(item_3))
        self.table.tableWidget.setItem(index, 4, QtWidgets.QTableWidgetItem(item_4))
        self.table.tableWidget.setItem(index, 5, QtWidgets.QTableWidgetItem(item_5))
        self.table.tableWidget.resizeRowToContents(index)
    
    def save_structure_row(self):               
        id_stru = self.table.tableWidget.currentIndex().row()
        
        if id_stru == -1:
            id_stru = 0
            
        self.con.open()
        
        try:
            title_text = self.table.tableWidget.currentItem().text()
        except:
            title_text = ""
        
        self.table.tableWidget.resizeRowToContents(id_stru)
        for nu, id_namb in enumerate(self.nambs):
            if id_namb >= id_stru:
                change = id_namb + 1
                self.nambs[nu] = change
                
                                        
                self.query.exec("UPDATE date SET id_1 = ? WHERE targetID = ?")                     
                self.query.addBindValue(change)
                self.query.addBindValue(nu)
                
                self.query.exec()
                
     
        self.nambs.append(id_stru)        
        self.query.prepare("INSERT INTO date VALUES (?, null, null, null, null, null, null, ?, null, ?, null, null)")
        self.query.addBindValue(id_stru)
        self.query.addBindValue(title_text)
        self.query.addBindValue(True) 
        self.query.exec()
        
        for nu, id_namb in enumerate(self.nambs):
            
            self.query.exec("UPDATE date SET targetID = ? WHERE id_1 = ?")                     
            self.query.addBindValue(nu)
            self.query.addBindValue(id_namb)
            
            self.query.exec()

        self.con.close()
        
        self.table.pushButton_2.blockSignals(False)
        self.table.pushButton_3.blockSignals(False)
        self.table.pushButton_8.clicked.disconnect(self.save_structure_row)
        self.table.pushButton_6.clicked.disconnect(self.delete_row)
        self.table.pushButton_6.clicked.connect(self.message_del)
        self.table.pushButton_12.clicked.disconnect(self.worningunfo)
        self.table.pushButton_12.clicked.connect(self.add_info)
        self.table.label_4.setText('<h1 style="color: rgb(73, 99, 250);">Структуру збережено...</h1>')
        self.table.label_4.setFont(QtGui.QFont("Times", 4, QtGui.QFont.Bold))
        self.table.label_4.show()
        self.get_date()
        self.time.singleShot(2000, lambda: self.table.label_4.setText(f'<h1 style="color: rgb(73, 99, 250);">{self.date_edit}</h1>'))
        self.update_json()
        
    def save_persondat_row (self):
        id_person = self.table.tableWidget.currentIndex().row()
        if id_person == -1:
            id_person = 0
              
        self.con.open()
        
        
        try:
            text2 = self.table.tableWidget.item(id_person, 0).text()
        except:
            text2 = ""
            
        try:
            text3 = self.table.tableWidget.item(id_person, 1).text()
        except:
            text3 = ""
            
        try:
            text4 = self.table.tableWidget.item(id_person, 2).text()
        except:
            text4 = ""
            
        try:
            text5 = self.table.tableWidget.item(id_person, 3).text()
        except:
            text5 = ""
            
        try:
            text6 = self.table.tableWidget.item(id_person, 4).text()
        except:
            text6 = ""
            
        try:
            text7 = self.table.tableWidget.item(id_person, 5).text()
        except:
            text7 = ""
        
                    
        self.table.tableWidget.resizeRowToContents(id_person)
        for nu, id_namb in enumerate(self.nambs):
            if id_namb >= id_person:
                change = id_namb + 1
                self.nambs[nu] = change
                                    
                self.query.exec("UPDATE date SET id_1 = ? WHERE targetID = ?")                     
                self.query.addBindValue(change)
                self.query.addBindValue(nu)
                
                self.query.exec()
                                                                                    
        self.nambs.append(id_person)        
        self.query.prepare("INSERT INTO date VALUES (?, ?, ?, ?, ?, ?, ?, null, null, ?, null, null)")
        self.query.addBindValue(id_person)
        self.query.addBindValue(text2)
        self.query.addBindValue(text3)
        self.query.addBindValue(text4)
        self.query.addBindValue(text5)
        self.query.addBindValue(text6)
        self.query.addBindValue(text7)
        self.query.addBindValue(False)
        self.query.exec()
        
        for nu, id_namb in enumerate(self.nambs):
            
            self.query.exec("UPDATE date SET targetID = ? WHERE id_1 = ?")                     
            self.query.addBindValue(nu)
            self.query.addBindValue(id_namb)
            
            self.query.exec()
            
        self.con.close()
               
        self.table.pushButton_2.blockSignals(False)
        self.table.pushButton_3.blockSignals(False)
        self.table.pushButton_8.clicked.disconnect(self.save_persondat_row)
        self.table.pushButton_6.clicked.disconnect(self.delete_row)
        self.table.pushButton_6.clicked.connect(self.message_del)
        self.table.pushButton_12.clicked.disconnect(self.worningunfo)
        self.table.pushButton_12.clicked.connect(self.add_info)
        self.table.label_4.setText('<h1 style="color: rgb(73, 99, 250);">Контакт збережено...</h1>')
        self.table.label_4.setFont(QtGui.QFont("Times", 4, QtGui.QFont.Bold))
        self.table.label_4.show()
        self.get_date()
        self.time.singleShot(2000, lambda: self.table.label_4.setText(f'<h1 style="color: rgb(73, 99, 250);">{self.date_edit}</h1>'))
        self.update_json() 
        
    def delete_row (self):
        
        id_clear_row = self.table.tableWidget.currentIndex().row()
        if id_clear_row == self.boost_row[0]:
            
            self.table.tableWidget.removeRow(id_clear_row)           
            self.table.pushButton_6.clicked.disconnect(self.delete_row)
            self.table.pushButton_6.clicked.connect(self.message_del)
            self.table.pushButton_12.clicked.disconnect(self.worningunfo)
            self.table.pushButton_12.clicked.connect(self.add_info)
            
            self.table.pushButton_2.blockSignals(False)
            self.table.pushButton_3.blockSignals(False)
            try:
                self.table.pushButton_8.clicked.disconnect(self.save_persondat_row)
            except:
                None
                
            try:
                self.table.pushButton_8.clicked.disconnect(self.save_structure_row)
            except:
                None
            
            self.table.label_4.setText('<h1 style="color: rgb(73, 99, 250);">Рядок видалено...</h1>')
            self.table.label_4.setFont(QtGui.QFont("Times", 4, QtGui.QFont.Bold))
            self.table.label_4.show()
            self.time.singleShot(2000, lambda: self.table.label_4.setText(f'<h1 style="color: rgb(73, 99, 250);">{self.date_edit}</h1>'))
    
    def message_del (self):
        
        id_cose = self.table.tableWidget.currentIndex().row()
        try:
            item_cose = self.table.tableWidget.item(id_cose, 0).text()
        except:
            item_cose = 'Порожній рядок'
        
        if id_cose != -1:
            self.msg.setWindowTitle("Видалити")
            self.msg.setText(f"Ви впевнені, що хочете видалити назву структуры/контакт: {item_cose}?")
            self.msg.setStandardButtons(QMessageBox.Ok|QMessageBox.Cancel)
            self.msg.buttonClicked.connect(self.delete)
            self.msg.exec_()
    
    def delete (self, choice):
        
        if choice.text() == 'OK':
            id_row = self.table.tableWidget.currentIndex().row()
            
            self.con.open()
            
            self.query.exec("SELECT namefile FROM date WHERE id_1 = ?")
            self.query.addBindValue(id_row)
            self.query.exec()
            if self.query.next():
                    boolbutns_value = self.query.value(0)
            
            file_name = f'docdat\{boolbutns_value}.docx'
            
            if os.path.isfile(file_name):
                os.remove(file_name)
            else:
                None
            
            self.table.tableWidget.removeRow(id_row)
                    
            self.query.prepare("DELETE FROM date WHERE id_1 = ?")
            self.query.addBindValue(id_row)
            self.query.exec()
            try:
                self.nambs.remove(id_row)
            except:
                None
            for nu_upd, id_namb_upd in enumerate(self.nambs):
                
                self.query.exec("UPDATE date SET targetID = ? WHERE id_1 = ?")                     
                self.query.addBindValue(nu_upd)
                self.query.addBindValue(id_namb_upd)
                
                self.query.exec()
                    
            for nu_del, id_namb_del in enumerate(self.nambs):
                if id_namb_del >= id_row:
                    change_del = id_namb_del - 1
                    self.nambs[nu_del] = change_del
                                                                   
                    self.query.exec("UPDATE date SET id_1 = ? WHERE targetID = ?")                     
                    self.query.addBindValue(change_del)
                    self.query.addBindValue(nu_del)
                    
                    self.query.exec()
                  
            self.con.close()
                
            self.msg.buttonClicked.disconnect(self.delete)
            self.table.label_4.setText('<h1 style="color: rgb(73, 99, 250);">Видалено...</h1>')
            self.table.label_4.setFont(QtGui.QFont("Times", 4, QtGui.QFont.Bold))
            self.table.label_4.show()
            self.get_date()
            self.time.singleShot(2000, lambda: self.table.label_4.setText(f'<h1 style="color: rgb(73, 99, 250);">{self.date_edit}</h1>'))
            self.update_json()
        
    def edit (self):
        self.con.open()
        
        findindex_1 = self.table.tableWidget.currentIndex().row()
        
        try:
            updateitem_1 = self.table.tableWidget.item(findindex_1, 0).text()
        except:
            updateitem_1 =''
        
        self.query.exec("UPDATE date SET name_structure = ? WHERE id_1 = ? AND Booleanegt = 1")                     
        self.query.addBindValue(updateitem_1)
        self.query.addBindValue(findindex_1)
        
        self.query.exec()
        
        try:
            updateitem_2 = self.table.tableWidget.item(findindex_1, 1).text()
        except:
            updateitem_2 =''
        
        try:
            updateitem_3 = self.table.tableWidget.item(findindex_1, 2).text()
        except:
            updateitem_3 =''
            
        try:
            updateitem_4 = self.table.tableWidget.item(findindex_1, 3).text()
        except:
            updateitem_4 =''
            
        try:
            updateitem_5 = self.table.tableWidget.item(findindex_1, 4).text()
        except:
            updateitem_5 =''
            
        try:
            updateitem_6 = self.table.tableWidget.item(findindex_1, 5).text()
        except:
            updateitem_6 =''
         
        self.query.exec("UPDATE date SET name = ?, position = ?, miniATC = ?, official_namber = ?, ATC_10 = ?, office = ? WHERE id_1 = ? AND Booleanegt = 0")                     
        self.query.addBindValue(updateitem_1)
        self.query.addBindValue(updateitem_2)
        self.query.addBindValue(updateitem_3)
        self.query.addBindValue(updateitem_4)
        self.query.addBindValue(updateitem_5)
        self.query.addBindValue(updateitem_6)
        self.query.addBindValue(findindex_1)
    
        self.query.exec()
        self.con.close()
        
        self.table.label_4.setText('<h1 style="color: rgb(73, 99, 250);">Змінено...</h1>')
        self.table.label_4.setFont(QtGui.QFont("Times", 4, QtGui.QFont.Bold))
        self.table.label_4.show()
        self.get_date()
        self.time.singleShot(2000, lambda: self.table.label_4.setText(f'<h1 style="color: rgb(73, 99, 250);">{self.date_edit}</h1>'))
        self.update_json()
        
    def slideLeftMenu(self):
        width = self.table.frame_2.width()
        
        if width == 0:
            
            newWidth = 200
            self.table.pushButton.setIcon(QtGui.QIcon(u'ico/right_arrow.png'))
            
        else:
            
            newWidth = 0
            self.table.pushButton.setIcon(QtGui.QIcon(u'ico/gamburg.png'))
            
        self.animation = QtCore.QPropertyAnimation(self.table.frame_2, b'maximumWidth')
        
        self.animation.setDuration(450)
        self.animation.setStartValue(width)
        self.animation.setEndValue(newWidth)
        self.animation.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
        self.animation.start()
    
    def slideSearchMenu(self, bool):
        
        if bool == True:
            width = self.table.frame_8.width()
            
            if width == 0:
                
                newWidth = 150
                
            else:
                
                newWidth = 0
                
            self.animation = QtCore.QPropertyAnimation(self.table.frame_8, b'maximumWidth')
            
            self.animation.setDuration(450)
            self.animation.setStartValue(width)
            self.animation.setEndValue(newWidth)
            self.animation.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation.start()
        
    def load_on_display(self):        
        self.con.open()
        self.query.exec("SELECT * FROM date")

        if self.query.isActive():
            self.query.first()
            
            while self.query.isValid():        
                try:
                    
                    self.nambs.append(self.query.value ('id_1'))
                except:
                    None
                self.query.next()

        self.query.exec("SELECT * FROM date ORDER BY id_1")

        if self.query.isActive():
            self.query.first()
                
            while self.query.isValid():                   
                if self.query.value ('Booleanegt') == True:         
                   
                    self.add_row_stru_on_display(self.query.value ('id_1'), self.query.value ('name_structure'))
                if self.query.value ('Booleanegt') == False:
                   
                    self.add_row_persondat_on_display(self.query.value ('id_1'), self.query.value ('name'), self.query.value ('position'), self.query.value ('miniATC'),self.query.value ('official_namber'), self.query.value ('ATC_10'), self.query.value ('office'))
                
                self.query.next()
        self.con.close()
    
    def number_of_finds_item(self):
        
        self.numb = 0
        
        for numb2 in self.testlist:
            self.numb += 1        
        self.table.label_3.setText(str(self.numb))
        self.table.label_3.setFont(QtGui.QFont("Times", 8, QtGui.QFont.Bold))
        self.dinamic_namb.insert(0, self.numb)
        
        if self.numb != 0:
            self.slideSearchMenu(self.check)
            self.check = False
        if self.numb == 0 and self.check == False:
            self.check = True
            self.slideSearchMenu(self.check)
            
    def arrow_right(self):
        
        cons = 1
        
        din_nam = self.dinamic_namb[0]
        minus = din_nam - cons
        
        
        if minus > -1:
            self.table.label_3.setText(str(minus + 1))
            try:
                colorbost = self.testlist[self.dinamic_namb[0]]
            except:
                None
                        
            self.dinamic_namb.insert(0, minus)
            viborca = self.testlist[self.dinamic_namb[0]]
            
            self.table.tableWidget.setCurrentItem(viborca, QtCore.QItemSelectionModel.Current)
            viborca.setBackground(QtGui.QColor(0, 109, 200, 255))
    
            try:
                self.table.tableWidget.setCurrentItem(colorbost.setBackground(QtGui.QColor(240, 240, 0, 127)))
            except:
                None
                  
    def arrow_left(self):
        cons = 1
        
        din_nam = self.dinamic_namb[0]
        plus = din_nam + cons
        
        
        if plus < self.numb:
            self.table.label_3.setText(str(plus + 1))
            try:
                colorbost = self.testlist[self.dinamic_namb[0]]
            except:
                None
                        
            self.dinamic_namb.insert(0, plus)
            viborca = self.testlist[self.dinamic_namb[0]]
            
            self.table.tableWidget.setCurrentItem(viborca, QtCore.QItemSelectionModel.Current)
            viborca.setBackground(QtGui.QColor(0, 109, 200, 255))
    
            try:
                self.table.tableWidget.setCurrentItem(colorbost.setBackground(QtGui.QColor(240, 240, 0, 127)))
            except:
                None
    
    def find_item(self, d):
    
        self.table.tableWidget.setCurrentItem(None)

        matching_items = self.table.tableWidget.findItems(d, QtCore.Qt.MatchContains)
        
        
        if not d:
            try:
                for item_full_clear in self.testlist:
                    item_full_clear.setBackground(QtGui.QColor(0, 0, 0, 0))
                self.testlist.clear()
                self.number_of_finds_item()
            except:
                None
            
            return
       
        for item in matching_items:
                    
         if item in self.testlist:

             try:
                for item2 in self.testlist:
                    item2.setBackground(QtGui.QColor(0, 0, 0, 0))
                self.testlist.clear()    
             except:
                None
       
        if matching_items:            

                for item in matching_items:
                    print (item)
                        
                    if item not in self.testlist:
                        
                        self.testlist.append(item)
                        
                        item.setBackground(QtGui.QColor(240, 240, 0, 127))
            
                    self.number_of_finds_item()
           
    def save_word(self):
        doc = docx.Document()    
        doc.add_heading('Телефонний довідник', 0)
        st = doc.styles['Normal']
        st.font.size = Pt(13)
        
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
                
        hdr_cells = table.rows[0].cells
        
        style_1 = hdr_cells[0].paragraphs[0]
        style_1.add_run('Прізвище, ім’я та по батькові').bold = True
        style_2 = hdr_cells[1].paragraphs[0]
        style_2.add_run('Посада').bold = True
        style_3 = hdr_cells[2].paragraphs[0]
        style_3.add_run('Міні – АТС внутрішній').bold = True
        style_4 = hdr_cells[3].paragraphs[0]
        style_4.add_run('Службовий міський').bold = True
        style_5 = hdr_cells[4].paragraphs[0]
        style_5.add_run('АТС-10').bold = True
        style_6 = hdr_cells[5].paragraphs[0]
        style_6.add_run('Кабінет').bold = True
        
        with open ('dict_tab.json', 'r') as filread:
            out_text = json.load(filread)
            
            for item in out_text:
                if item ['Booleanegt'] == True:
                      table = doc.add_table(rows=1, cols=1)
                      cell = table.rows[0].cells
                      table.style = 'Table Grid'
                      sty = cell[0].paragraphs[0]
                      sty.alignment = WD_ALIGN_PARAGRAPH.CENTER
                      sty.add_run(item ['name_structure']).bold = True
                      
                if item ['Booleanegt'] == False:
                      table = doc.add_table(rows=1, cols=6)
                      table.style = 'Table Grid'
                      name_cell = table.rows[0].cells
                      name_cell[0].text = item ['name']
                      name_cell[1].text = item ['position']
                      name_cell[2].text = item ['miniATC']
                      name_cell[3].text = item ['official_namber']
                      name_cell[4].text = item ['ATC_10']
                      name_cell[5].text = item ['office']
        try:              
            SaveFileDiolog = QtWidgets.QFileDialog.getSaveFileName(caption = 'Save file', directory = 'телефонний довідник.docx', filter = 'Save file (*.docx)')

            way = []
            for data_way in SaveFileDiolog:
                way.append(data_way)
            
            doc.save(way[0])
            
        except:
            None
    
    def update_json (self):
        self.update_data_list.clear()      
        self.con.open()
        self.query.exec("SELECT * FROM date ORDER BY id_1")
        
        if self.query.isActive():
            self.query.first()
                
            while self.query.isValid():
                boole = self.query.value ('Booleanegt') 
                Id = self.query.value ('id_1')
                name = self.query.value ('name')
                position = self.query.value ('position')           
                miniATC = self.query.value ('miniATC') 
                official_namber = self.query.value ('official_namber')
                ATC_10 = self.query.value ('ATC_10')
                office = self.query.value ('office')
                name_structure = self.query.value ('name_structure')
                targetID = self.query.value ('targetID')
                Boolbutns = self.query.value ('Boolbutns')
                namefile = self.query.value ('namefile')
                dict_data_table = {
                                   'Booleanegt': boole,
                                   'id_1': Id,
                                   'name': name,
                                   'position': position,
                                   'miniATC': miniATC,
                                   'official_namber': official_namber,
                                   'ATC_10': ATC_10,
                                   'office': office,
                                   'name_structure': name_structure,
                                   'targetID': targetID,
                                   'Boolbutns': Boolbutns,
                                   'namefile': namefile,
                                   'date': self.date_edit
                                   }
                self.update_data_list.append(dict_data_table)
                self.query.next()
           
            with open ('docdat\\dict_tab.json', "w", encoding='utf-8') as this_w:
                json.dump(self.update_data_list, this_w, indent=1)
            self.con.close()
    def generate_unique_string(self, length):
        return ''.join(random.choices(string.ascii_letters + string.digits, k=length))
    def add_info (self):
        
        if self.table.tableWidget.selectedItems():
            id_ro = self.table.tableWidget.currentIndex().row()
            self.con.open()
            self.query.exec("SELECT Boolbutns FROM date WHERE id_1 = ?")
            self.query.addBindValue(id_ro)
            self.query.exec()
            if self.query.next():
                    boolbutns_value = self.query.value(0)  # Получить значение из столбца Boolbutns
            self.con.close()
            print(boolbutns_value)
            
            if boolbutns_value != 1:
                self.msg.setWindowTitle("Увага")
                self.msg.setText("Додати інформацію до вибраного контакту/структури ?")
                self.msg.setStandardButtons(QMessageBox.Ok|QMessageBox.Cancel)
                result = self.msg.exec_()
                
                if result == QMessageBox.Ok:
                    new_name = self.generate_unique_string(10)
                    doc = Document()
                    doc.add_paragraph('Привет, мир!')
                    file_path = f'docdat\\{new_name}.docx'
                    doc.save(file_path)
                    
                    icon = QtGui.QIcon()
                    self.nativebutton = QtWidgets.QPushButton()
                    sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Minimum)
                    sizePolicy.setHorizontalStretch(0)
                    sizePolicy.setVerticalStretch(0)
                    sizePolicy.setHeightForWidth(self.nativebutton.sizePolicy().hasHeightForWidth())
                    self.nativebutton.setMinimumSize(QtCore.QSize(25, 25))
                    self.nativebutton.setMaximumSize(QtCore.QSize(25, 25))
                    self.nativebutton.setStyleSheet("QPushButton{\n"
            "    border: none\n"
            "}\n"
            "\n"
            "QPushButton {\n"
            "    border-radius: 5px\n"
            "}\n"
            "\n"
            "QPushButton:hover{\n"
            "    background-color:rgb(41, 198, 255)\n"
            "}")
                    icon.addPixmap(QtGui.QPixmap("ico/info.png"), QtGui.QIcon.Normal, QtGui.QIcon.Off)
                    self.nativebutton.setIcon(icon)
                    self.nativebutton.setIconSize(QtCore.QSize(25, 25))
                    
                    self.table.tableWidget.setCellWidget(id_ro, 6, self.nativebutton)
                    
                    self.con.open()
                    
                    self.query.exec("UPDATE date SET Boolbutns = ?, namefile = ? WHERE id_1 = ?")                     
                    self.query.addBindValue(True)
                    self.query.addBindValue(new_name)
                    self.query.addBindValue(id_ro)            
                    self.query.exec()
                    
                    self.con.close()
                    self.update_json()
                    print (id_ro)
            else:
                self.msg.setWindowTitle("Увага")
                self.msg.setText("У вибраному контакті/структурі інформація вже додана.")
                self.msg.setStandardButtons(QMessageBox.Ok)
                self.msg.exec_()
        else:
            self.msg.setWindowTitle("Увага")
            self.msg.setText("Оберіть структуру або контакт до якого потрібно додати інформацію або спочатку збережіть нові дані контакта/структури.")
            self.msg.setStandardButtons(QMessageBox.Ok)
            self.msg.exec_()
        self.nativebutton.clicked.connect(self.natifBatnfunc)
    
    def natifBatnfunc (self):
        id_ro = self.table.tableWidget.currentIndex().row()
        self.con.open()
        self.query.exec("SELECT namefile FROM date WHERE id_1 = ?")
        self.query.addBindValue(id_ro)
        self.query.exec()
        if self.query.next():
                boolbutns_value = self.query.value(0)  
        self.con.close()
        file_name = f'docdat\\{boolbutns_value}.docx'
        
        if os.path.isfile(file_name):         
            os.startfile(file_name)
        else:
            self.msg.setWindowTitle("Увага")
            self.msg.setText("Документ з інформацією зазначенї структури/контакту відстуній. Створено новий документ.")
            self.msg.setStandardButtons(QMessageBox.Ok)
            self.msg.exec_()
            doc = Document()
            doc.add_paragraph('Привет, мир!')
            doc.save(file_name)
            os.startfile(file_name)
        print ("нажал нативную в строке №", id_ro, "название документа =>", boolbutns_value)
    
    def addTitleInfo (self):
        file_path = 'docdat\\titleinfo.docx'
        if os.path.isfile(file_path):         
            os.startfile(file_path)
        else:        
            doc = Document()
            doc.add_paragraph('Тут повинно бути положення')
            file_path = 'docdat\\titleinfo.docx'
            doc.save(file_path)
            os.startfile(file_path)
            
    def posButnonDisplay (self):
        
        self.con.open()
        self.query.exec("SELECT * FROM date")

        if self.query.isActive():
            self.query.first()
            
            while self.query.isValid():
                if self.query.value ('Boolbutns') == True:
                    
                   icon = QtGui.QIcon()
                   self.nativebutton = QtWidgets.QPushButton()
                   sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Minimum)
                   sizePolicy.setHorizontalStretch(0)
                   sizePolicy.setVerticalStretch(0)
                   sizePolicy.setHeightForWidth(self.nativebutton.sizePolicy().hasHeightForWidth())
                   self.nativebutton.setMinimumSize(QtCore.QSize(25, 25))
                   self.nativebutton.setMaximumSize(QtCore.QSize(25, 25))
                   self.nativebutton.setStyleSheet("QPushButton{\n"
           "    border: none\n"
           "}\n"
           "\n"
           "QPushButton {\n"
           "    border-radius: 5px\n"
           "}\n"
           "\n"
           "QPushButton:hover{\n"
           "    background-color:rgb(41, 198, 255)\n"
           "}")
                   icon.addPixmap(QtGui.QPixmap("ico/info.png"), QtGui.QIcon.Normal, QtGui.QIcon.Off)
                   self.nativebutton.setIcon(icon)
                   self.nativebutton.setIconSize(QtCore.QSize(25, 25))
        
                   self.table.tableWidget.setCellWidget(self.query.value ('id_1'), 6, self.nativebutton)
                   self.nativebutton.clicked.connect(self.natifBatnfunc)
                self.query.next()

        self.con.close()
    
    def worningunfo(self):
        self.msg.setWindowTitle("Увага")
        self.msg.setText("Cпочатку збережіть нові дані контакта/структури.")
        self.msg.setStandardButtons(QMessageBox.Ok)
        self.msg.exec_()
    
    
    def start_server(self):
        os.startfile('telephone_directory_server.exe')

if __name__ == '__main__':
    app = QtWidgets.QApplication(sys.argv)  
    window = TableBook()
    window.setWindowTitle('Телефонний довідник')
    window.setWindowIcon(QtGui.QIcon('ico/til.png'))
    window.Sql()
    window.buttons()
    sys.exit(app.exec_())
