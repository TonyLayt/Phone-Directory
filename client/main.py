# -*- coding: utf-8 -*-
"""
Created on Wed Dec 28 16:46:13 2022

@author: m.s.kuznietsov
"""

from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtWidgets import QMessageBox
from uitabl import Ui_MainWindow
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm
import configparser
import socket
import json
import sys
import os



class TableBook (QtWidgets.QMainWindow):
    def __init__(self):
        self.nambs = []
        QtWidgets.QMainWindow.__init__(self)
        self.table = Ui_MainWindow()
        self.table.setupUi(self)        
        self.show()
        self.config = configparser.ConfigParser()
        self.ini_config()
        self.ip = str(self.config['Connection Settings']['IP-адреса(TCP/IPv4)'])
        self.port = int(self.config['Connection Settings']['Port'])
        self.testlist = []
        self.dinamic_namb = []
        self.check = True
         
    def buttons (self):
        self.table.pushButton.clicked.connect(self.slideLeftMenu)
        self.table.lineEdit.textChanged.connect(self.find_item)
        self.table.pushButton_3.clicked.connect(self.connect_to_server)
        self.table.pushButton_5.clicked.connect(self.save_word)
        self.table.pushButton_9.clicked.connect(self.titleInfo)
        self.table.pushButton_10.clicked.connect(self.arrow_left)
        self.table.pushButton_11.clicked.connect(self.arrow_right)
        
    
    def ini_config (self):
        if not self.config.read('Settings_connect.ini'):
            self.config['Connection Settings'] = {'IP-адреса(TCP/IPv4)':'localhost', 'Port':'8000'}
            with open ('Settings_server.ini', 'w') as configfile:
                self.config.write(configfile)
        
    def add_row_stru_on_display (self, index, item): 
        self.table.tableWidget.insertRow(index)
        self.table.tableWidget.setSpan (index, 0, 1, 6)
        
        styl_text_load_structure = QtWidgets.QTableWidgetItem()
        styl_text_load_structure.setTextAlignment(QtCore.Qt.AlignCenter)
        
        font_text_load_structure = QtGui.QFont()
        font_text_load_structure.setPointSize(12)
        font_text_load_structure.setBold(True)
        font_text_load_structure.setWeight(75)
        styl_text_load_structure.setFont(font_text_load_structure)
        self.table.tableWidget.setItem(index, 0, styl_text_load_structure)
        item_header = self.table.tableWidget.item(index, 0)
        item_header.setText(item)
        self.table.tableWidget.resizeRowToContents(index)   
        
    def add_row_persondat_on_display (self, index, item_0, item_1, item_2, item_3, item_4, item_5): 
        self.table.tableWidget.insertRow(index)
        self.table.tableWidget.setItem(index, 0, QtWidgets.QTableWidgetItem(item_0))
        self.table.tableWidget.setItem(index, 1, QtWidgets.QTableWidgetItem(item_1))
        self.table.tableWidget.setItem(index, 2, QtWidgets.QTableWidgetItem(item_2))
        self.table.tableWidget.setItem(index, 3, QtWidgets.QTableWidgetItem(item_3))
        self.table.tableWidget.setItem(index, 4, QtWidgets.QTableWidgetItem(item_4))
        self.table.tableWidget.setItem(index, 5, QtWidgets.QTableWidgetItem(item_5))
        self.table.tableWidget.resizeRowToContents(index)
        
    def slideLeftMenu(self): 
        width = self.table.frame_2.width()
        
        if width == 0:
            
            newWidth = 200
            self.table.pushButton.setIcon(QtGui.QIcon(u'ico/right_arrow.png'))
            
        else:
            
            newWidth = 0
            self.table.pushButton.setIcon(QtGui.QIcon(u'ico/gamburg.png'))
            
        self.animation = QtCore.QPropertyAnimation(self.table.frame_2, b'maximumWidth')
        
        self.animation.setDuration(450)
        self.animation.setStartValue(width)
        self.animation.setEndValue(newWidth)
        self.animation.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
        self.animation.start()
    
    def slideSearchMenu(self, bool):
        
        if bool == True:
            width = self.table.frame_8.width()
            
            if width == 0:
                
                newWidth = 150
                
            else:               
                newWidth = 0
                
            self.animation = QtCore.QPropertyAnimation(self.table.frame_8, b'maximumWidth')
            
            self.animation.setDuration(450)
            self.animation.setStartValue(width)
            self.animation.setEndValue(newWidth)
            self.animation.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation.start()
    
    def number_of_finds_item(self):
        
        self.numb = 0
        
        for numb2 in self.testlist:
            self.numb += 1
       
        self.table.label_4.setText(str(self.numb))
        self.table.label_4.setFont(QtGui.QFont("Times", 8, QtGui.QFont.Bold))
        self.dinamic_namb.insert(0, self.numb)
        
        if self.numb != 0:
            
            self.slideSearchMenu(self.check)
            self.check = False
        if self.numb == 0 and self.check == False:
            
            self.check = True
            self.slideSearchMenu(self.check)
            
    def arrow_right(self):
        
        cons = 1
        
        din_nam = self.dinamic_namb[0]
        minus = din_nam - cons
        
        
        if minus > -1:
            self.table.label_4.setText(str(minus + 1))
            try:
                colorbost = self.testlist[self.dinamic_namb[0]]
            except:
                None
                        
            self.dinamic_namb.insert(0, minus)
            viborca = self.testlist[self.dinamic_namb[0]]
            
            self.table.tableWidget.setCurrentItem(viborca, QtCore.QItemSelectionModel.Current)
            viborca.setBackground(QtGui.QColor(0, 109, 200, 255))
    
            try:
                self.table.tableWidget.setCurrentItem(colorbost.setBackground(QtGui.QColor(240, 240, 0, 127)))
            except:
                None
              
    
    def arrow_left(self):
        
        cons = 1
        
        din_nam = self.dinamic_namb[0]
        plus = din_nam + cons
        
        
        if plus < self.numb:
            self.table.label_4.setText(str(plus + 1))
            try:
                colorbost = self.testlist[self.dinamic_namb[0]]
            except:
                None
                        
            self.dinamic_namb.insert(0, plus)
            viborca = self.testlist[self.dinamic_namb[0]]
            
            self.table.tableWidget.setCurrentItem(viborca, QtCore.QItemSelectionModel.Current)
            viborca.setBackground(QtGui.QColor(0, 109, 200, 255))
    
            try:
                self.table.tableWidget.setCurrentItem(colorbost.setBackground(QtGui.QColor(240, 240, 0, 127)))
            except:
                None
    
    def find_item(self, d):
        
        self.table.tableWidget.setCurrentItem(None)  
        matching_items = self.table.tableWidget.findItems(d, QtCore.Qt.MatchContains)
           
        if not d:
            try:
                for item_full_clear in self.testlist:
                    item_full_clear.setBackground(QtGui.QColor(0, 0, 0, 0))
                self.testlist.clear()
                self.number_of_finds_item()
            except:
                None           
            return
              
        for item in matching_items:                   
         if item in self.testlist:
             try:
                for item2 in self.testlist:
                    item2.setBackground(QtGui.QColor(0, 0, 0, 0))
                self.testlist.clear()    
             except:
                None
                                  
        if matching_items:                                      
                for item in matching_items:
                    
                        
                    if item not in self.testlist:
                        
                        self.testlist.append(item)
                        
                        item.setBackground(QtGui.QColor(240, 240, 0, 127))
            
                    self.number_of_finds_item()
        
    def connect_to_server(self):
        try:
           user_socket = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
           user_socket.connect((self.ip, self.port))
           self.table.tableWidget.setRowCount(0)
           
           save_folder = 'docdat'
           
           num_files = int(user_socket.recv(1024).decode())
           

           for _ in range(num_files):
               # Получаем имя файла и его размер от сервера
               file_name = user_socket.recv(1024).decode()
               file_size = int(user_socket.recv(1024).decode())
        
               file_path = os.path.join(save_folder, file_name)
        
               # Получаем содержимое файла от сервера и сохраняем его
               with open(file_path, 'wb') as file:
                   total_bytes_received = 0
                   while total_bytes_received < file_size:
                       file_data = user_socket.recv(1024)
                       total_bytes_received += len(file_data)
                       file.write(file_data)
           user_socket.close()
           self.date_load()
        except:
            self.table.label_3.setText('<h1 style="color: rgb(198, 22, 22);">З`єднання відсутнє...</h1>')
            self.table.label_3.setFont(QtGui.QFont("Times", 4, QtGui.QFont.Bold))
            
            
    
    def save_word(self):
        doc = docx.Document()
        doc.add_picture('ico/gerb2.png', width=Mm(25))
        doc.add_heading('Телефонний довідник Мінсоцполітики', 0)
        st = doc.styles['Normal']
        st.font.size = Pt(13)
        
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
                
        hdr_cells = table.rows[0].cells
        
        style_1 = hdr_cells[0].paragraphs[0]
        style_1.add_run('Прізвище, ім’я та по батькові').bold = True
        style_2 = hdr_cells[1].paragraphs[0]
        style_2.add_run('Посада').bold = True
        style_3 = hdr_cells[2].paragraphs[0]
        style_3.add_run('Міні – АТС внутрішній').bold = True
        style_4 = hdr_cells[3].paragraphs[0]
        style_4.add_run('Службовий міський').bold = True
        style_5 = hdr_cells[4].paragraphs[0]
        style_5.add_run('АТС-10').bold = True
        style_6 = hdr_cells[5].paragraphs[0]
        style_6.add_run('Кабінет').bold = True
        
        with open ('docdat\\dict_tab.json', 'r') as filread:
            out_text = json.load(filread)
            
            for item in out_text:
                if item ['Booleanegt'] == True:
                      table = doc.add_table(rows=1, cols=1)
                      cell = table.rows[0].cells
                      table.style = 'Table Grid'
                      sty = cell[0].paragraphs[0]
                      sty.alignment = WD_ALIGN_PARAGRAPH.CENTER
                      sty.add_run(item ['name_structure']).bold = True
                      
                if item ['Booleanegt'] == False:
                      table = doc.add_table(rows=1, cols=6)
                      table.style = 'Table Grid'
                      name_cell = table.rows[0].cells
                      name_cell[0].text = item ['name']
                      name_cell[1].text = item ['position']
                      name_cell[2].text = item ['miniATC']
                      name_cell[3].text = item ['official_namber']
                      name_cell[4].text = item ['ATC_10']
                      name_cell[5].text = item ['office']
        try:              
            SaveFileDiolog = QtWidgets.QFileDialog.getSaveFileName(caption = 'Save file', directory = 'телефонний довідник.docx', filter = 'Save file (*.docx)')

            way = []
            for data_way in SaveFileDiolog:
                way.append(data_way)
            
            doc.save(way[0])
            
        except:
            None

    
    def date_load(self):
        try:
            with open ('docdat\\dict_tab.json', 'r') as filread:
                out_text = json.load(filread)
                
                for item in out_text:
                    if item ['Booleanegt'] == True:         
                         
                          self.add_row_stru_on_display(item ['id_1'], item ['name_structure'])
         
                    if item ['Booleanegt'] == False:
                          
                          self.add_row_persondat_on_display(item ['id_1'], item ['name'], item ['position'], item ['miniATC'], item ['official_namber'], item ['ATC_10'], item ['office'])
                          
                    if item ['Boolbutns'] == True:
                        
                        icon = QtGui.QIcon()
                        self.nativebutton = QtWidgets.QPushButton()
                        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Minimum)
                        sizePolicy.setHorizontalStretch(0)
                        sizePolicy.setVerticalStretch(0)
                        sizePolicy.setHeightForWidth(self.nativebutton.sizePolicy().hasHeightForWidth())
                        self.nativebutton.setMinimumSize(QtCore.QSize(25, 25))
                        self.nativebutton.setMaximumSize(QtCore.QSize(25, 25))
                        self.nativebutton.setStyleSheet("QPushButton{\n"
                "    border: none\n"
                "}\n"
                "\n"
                "QPushButton {\n"
                "    border-radius: 5px\n"
                "}\n"
                "\n"
                "QPushButton:hover{\n"
                "    background-color:rgb(41, 198, 255)\n"
                "}")
                        icon.addPixmap(QtGui.QPixmap("ico/info.png"), QtGui.QIcon.Normal, QtGui.QIcon.Off)
                        self.nativebutton.setIcon(icon)
                        self.nativebutton.setIconSize(QtCore.QSize(25, 25))
             
                        self.table.tableWidget.setCellWidget(item ['id_1'], 6, self.nativebutton)
                        self.nativebutton.clicked.connect(self.natifBatnfunc)
                        
                self.table.label_3.setText(f'<h1 style="color: rgb(73, 99, 250);">{item["date"]}</h1>')
                self.table.label_3.setFont(QtGui.QFont("Times", 4, QtGui.QFont.Bold))       
                          
        except:
             None

    def natifBatnfunc (self):
        print ('нажал')
        id_ro = self.table.tableWidget.currentIndex().row()
        
        try:
            with open ('docdat\\dict_tab.json', 'r') as filread:
                out_text = json.load(filread)
                
                for item in out_text:
                    
                    if item ['id_1'] == id_ro and item ['Boolbutns'] == True:
                        
                        file_name = item ['namefile']
                       
                        if os.path.isfile(f'docdat\\{file_name}.docx'):
                            
                            os.startfile(f'docdat\\{file_name}.docx')
                        
                        else:
                            msg = QMessageBox()
                            msg.setIcon(QMessageBox.Warning)
                            msg.setWindowTitle("Увага")
                            msg.setText("Данні не знайдено, оновіть інформацію або зверніться до адміністратора.")
                            msg.setStandardButtons(QMessageBox.Ok)
                            msg.exec_()
                        
        except:
             msg = QMessageBox()
             msg.setIcon(QMessageBox.Warning)
             msg.setWindowTitle("Увага")
             msg.setText("Помилка при відкритті файла, оновіть інформацію або зверніться до адміністратора.")
             msg.setStandardButtons(QMessageBox.Ok)
             msg.exec_()
    
    def titleInfo (self):
        file_path = 'docdat\\titleinfo.docx'
        if os.path.isfile(file_path):         
            os.startfile(file_path)
        else:    
            msg = QMessageBox()
            msg.setIcon(QMessageBox.Warning)
            msg.setWindowTitle("Увага")
            msg.setText("Данні не знайдено, оновіть інформацію або зверніться до адміністратора.")
            msg.setStandardButtons(QMessageBox.Ok)
            msg.exec_()
            
       
if __name__ == '__main__':
    app = QtWidgets.QApplication(sys.argv)  
    window = TableBook()
    window.setWindowTitle('Телефонний довідник')
    window.setWindowIcon(QtGui.QIcon('ico/til.png'))
    window.buttons()
    window.date_load()
    sys.exit(app.exec_())
